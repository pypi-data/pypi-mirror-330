import datetime
import os
import argparse
from docx import Document
from docx.shared import Inches, Pt
from tqdm import tqdm
from PIL import Image
import traceback


def get_images_from_folder(folder):
    """Returns a list of image file paths from the given folder, including .thumb files."""
    valid_extensions = (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".thumb")
    return [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(valid_extensions)]


def add_image_keeping_ratio(run, img_path, max_width, max_height):
    """Adds an image to a Word document while maintaining aspect ratio."""
    with Image.open(img_path) as img:
        width, height = img.size
        dpi = img.info.get("dpi", (72, 72))  # Default DPI to 72 if missing
        width_in_inches = width / dpi[0]
        height_in_inches = height / dpi[1]

        ratio = min(max_width / width_in_inches, max_height / height_in_inches)
        new_width = width_in_inches * ratio
        new_height = height_in_inches * ratio

        run.add_picture(img_path, width=Inches(new_width), height=Inches(new_height))


import traceback


from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx_from_images(image_paths, rows_per_page, output_doc, max_width, max_height):
    """Generates a DOCX file with a specified number of images per page, keeping aspect ratio."""
    try:
        if not image_paths:
            raise ValueError("No images found in the current directory.")

        doc = Document()
        total_images = len(image_paths)
        cols_per_row = 5
        error_list = []

        for i in tqdm(range(0, total_images, rows_per_page * cols_per_row), desc="Processing Images"):
            page_images = image_paths[i : i + rows_per_page * cols_per_row]
            table = doc.add_table(rows=rows_per_page, cols=cols_per_row)

            row_idx, col_idx = 0, 0
            for img_path in page_images:
                try:
                    cell = table.cell(row_idx, col_idx)
                    cell_paragraph = cell.paragraphs[0]
                    run = cell_paragraph.add_run()
                    add_image_keeping_ratio(run, img_path, max_width, max_height)

                    file_name = os.path.basename(img_path)
                    file_name_row = cell_paragraph.add_run(f"\n{file_name}")
                    file_name_row.font.bold = True
                    file_name_row.font.size = Pt(5)

                    # Center text and image
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    col_idx += 1
                    if col_idx == cols_per_row:
                        col_idx = 0
                        row_idx += 1
                        if row_idx >= rows_per_page:
                            break
                except Exception as e:
                    error_msg = f"File: {img_path}: {e}"
                    print(error_msg)
                    error_list.append(error_msg)

            if i + rows_per_page * cols_per_row < total_images:
                doc.add_page_break()

        doc.save(output_doc)
        if error_list:
            print("Some images could not be processed:")
            for err in error_list:
                print(f" - {err}")

        print(f"DOCX file created successfully: {output_doc}")

    except ValueError as ve:
        print(f"ERROR: {ve}")
    except FileNotFoundError as fnfe:
        print(f"File not found: {fnfe}")
    except Exception as ex:
        print(f"Unexpected error: {ex}")
        traceback.print_exc()  # Prints the full traceback for better debugging


def main():
    """Main function to find images and create a DOCX file."""
    parser = argparse.ArgumentParser(description="Generate DOCX from images.")
    parser.add_argument("-r", "--rows_per_page", type=int, default=5, help="Number of rows per page in the DOCX.")
    parser.add_argument("-s", "--scale", type=float, default=1.3, help="Max image width and height in inches while preserving aspect ratio.")
    parser.add_argument("-o", "--output", type=str, default=f"combined_images_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx", help="Output DOCX file name.")
    parser.add_argument("-l", "--limit_images", type=int, default=0, help="Limit the number of images to process. Use 0 for no limit.")
    args = parser.parse_args()

    current_folder = os.getcwd()
    images = get_images_from_folder(current_folder)

    if args.limit_images != 0:
        images = images[: args.limit_images]

    if images:
        create_docx_from_images(images, args.rows_per_page, args.output, args.scale, args.scale)
    else:
        print("No image files found in the current folder.")


if __name__ == "__main__":
    main()
