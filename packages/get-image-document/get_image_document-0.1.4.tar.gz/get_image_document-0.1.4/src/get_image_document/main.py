import datetime
import os
import argparse
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from PIL import Image
from docx.shared import Pt


def get_images_from_folder(folder):
    """Returns a list of image file paths from the given folder, including .thumb files."""
    valid_extensions = (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".thumb")
    return [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(valid_extensions)]


def add_image_keeping_ratio(run, img_path, max_width, max_height):
    """Adds an image to a Word document while maintaining its original aspect ratio."""
    with Image.open(img_path) as img:
        width, height = img.size
        ratio = min(max_width / width, max_height / height)
        new_width = width * ratio
        new_height = height * ratio
        run.add_picture(img_path, width=Inches(new_width), height=Inches(new_height))


def create_docx_from_images(image_paths, rows_per_page, output_doc, max_width, max_height):
    """Generates a DOCX file with a specified number of images per page, keeping aspect ratio."""
    if not image_paths:
        print("No images found in the current directory.")
        return

    doc = Document()
    total_images = len(image_paths)

    for i in tqdm(range(0, total_images, rows_per_page * 5), desc="Processing Images"):  # 5 columns per row
        page_images = image_paths[i : i + rows_per_page * 5]
        table = doc.add_table(rows=rows_per_page, cols=5)

        row_idx, col_idx = 0, 0
        for img_path in page_images:
            if os.path.exists(img_path):
                cell = table.cell(row_idx, col_idx)
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.add_run()

                add_image_keeping_ratio(run, img_path, max_width, max_height)

                file_name = os.path.basename(img_path)
                file_name_row = cell_paragraph.add_run(f"\n{file_name}")
                file_name_row.font.bold = True
                file_name_row.font.size = Pt(5)

                # Centering text and image
                cell.paragraphs[0].alignment = 1
                for p in cell.paragraphs:
                    p.alignment = 1

            col_idx += 1
            if col_idx == 5:
                col_idx = 0
                row_idx += 1

        doc.add_page_break()

    doc.save(output_doc)
    print(f"DOCX file created successfully: {output_doc}")


def main():
    """Main function to find images and create a DOCX file."""
    parser = argparse.ArgumentParser(description="Generate DOCX from images.")
    parser.add_argument("-r", "--rows_per_page", type=int, default=5, help="Number of rows per page in the DOCX.")
    parser.add_argument("-s", "--scale", type=float, default=1.3, help="Max width and height of the image (inch)")
    parser.add_argument("-o", "--output", type=str, default=f"combined_images_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx", help="Output DOCX file name.")
    parser.add_argument("-l", "--limit_images", type=int, default=0, help="Limit the number of images to process. Use 0 for no limit.")
    args = parser.parse_args()

    current_folder = os.getcwd()
    images = get_images_from_folder(current_folder)

    if args.limit_images != 0:
        images = images[: args.limit_images]

    if images:
        create_docx_from_images(images, args.rows_per_page, args.output, args.scale, args.scale)
    else:
        print("No image files found in the current folder.")


if __name__ == "__main__":
    main()
