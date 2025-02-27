"""
Interesting links:
- https://stackoverflow.com/questions/61801936/set-background-color-shading-on-normal-text-with-python-docx
- https://github.com/ArtifexSoftware/pdf2docx/issues/54#issuecomment-715925252
"""
from yta_general_utils.programming.output import Output
from yta_general_utils.file.enums import FileExtension
from yta_general_utils.programming.validator.parameter import ParameterValidator
from docx.document import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from typing import Union

import io


class DocxCreator:
    """
    Class to simplify the creation of a DOCX.
    """

    @staticmethod
    def add_header(
        docx: Document,
        text: str,
        level: int = 1
    ):
        """
        Add a heading 'text' within the given 'level'
        (that must be between 0 and 9).
        """
        docx.add_heading(text, level = level)

    @staticmethod
    def add_paragraph(
        docx: Document,
        text: str
    ):
        """
        Add a simple paragraph.
        """
        docx.add_paragraph(text)

    @staticmethod
    def add_image(
        docx: Document,
        filename: str
    ):
        """
        Add an image to the document.
        """
        docx.add_picture(filename)

    @staticmethod
    def add_centered_image(
        docx: Document,
        filename: str
    ):
        """
        Add an image to the document, but centered.
        """
        # TODO: This is being tested
        tabla = docx.add_table(rows = 1, cols = 1)
        celda = tabla.cell(0, 0)
        celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        celda.paragraphs[0].add_run().add_picture(filename, width  = Inches(4))

    @staticmethod
    def get_bytearray(
        docx: Document
    ) -> bytearray:
        """
        Get the 'docx' bytearray as if it was
        written to a file but without doing it.
        """
        ParameterValidator.validate_mandatory_instance_of('docx', docx, Document)

        out = io.BytesIO()
        docx.save(out)

        return out.getvalue()

    @staticmethod
    def write(
        docx: Document,
        output_filename: Union[str, None] = None
    ) -> str:
        """
        Create the provided 'docx' as a file and
        return the filename with which the file
        has been stored locally.
        """
        ParameterValidator.validate_mandatory_instance_of('docx', docx, Document)

        output_filename = Output.get_filename(output_filename, FileExtension.DOCX)
        # TODO: I could return the docx bytearray instead
        # (check 'get_bytearray' method) (?)
        docx_bytearray = docx.save(output_filename)

        return output_filename