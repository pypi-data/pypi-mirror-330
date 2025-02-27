"""
Module to hold some inspirational examples.

Tutorial here:
- https://python-docx.readthedocs.io/en/latest/
"""
from dal_docx_creator.creator import DocxCreator
from docx import Document


__all__ = [
    'DocxExample'
]

class DocxExample:
    """
    Class to wrap some example of DOCXs.
    """

    @staticmethod
    def test_one():
        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        #document.add_picture('monty-truth.png', width=Inches(1.25))

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        document.add_page_break()

        DocxCreator.write(document, 'test_simple.docx')

    def test_paragraph_plus_centered_image():
        document = Document()
        DocxCreator.add_header(document, 'Ejemplo de prueba')
        DocxCreator.add_paragraph(document, 'Esto es un texto largo que habré intentado que se invente ChatGPT, blablaba...')
        c = 'C:/Users/dania/Desktop/3dgreenscreen.jpg'
        DocxCreator.add_centered_image(document, c)
        DocxCreator.write(document, 'test_parimg.docx')